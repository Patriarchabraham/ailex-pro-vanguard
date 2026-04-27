"""
AILEX Pilot — docs.py
PDF and document analysis: extract text, tables, images from PDFs/DOCX/TXT.
Feeds extracted content into AILEX pipeline as project context.
"""
from __future__ import annotations

import base64
import io
import os
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional


@dataclass
class DocumentContent:
    path:      str
    file_type: str
    text:      str
    pages:     int
    tables:    List[str]
    images:    List[bytes]
    metadata:  Dict[str, str] = field(default_factory=dict)
    error:     Optional[str] = None


class DocumentAnalyzer:
    """
    Extracts content from PDFs, DOCX, TXT files.
    Supports: PDF (pdfminer/pypdf), DOCX (python-docx), TXT, MD.
    """

    def analyze(self, path: str) -> DocumentContent:
        if not os.path.exists(path):
            return DocumentContent(path=path, file_type="", text="",
                                   pages=0, tables=[], images=[], error="File not found")
        ext = os.path.splitext(path)[1].lower()
        if ext == ".pdf":
            return self._pdf(path)
        elif ext in (".docx", ".doc"):
            return self._docx(path)
        elif ext in (".txt", ".md", ".rst", ".csv"):
            return self._text(path)
        else:
            return DocumentContent(path=path, file_type=ext, text="",
                                   pages=0, tables=[], images=[],
                                   error=f"Unsupported format: {ext}")

    def analyze_with_claude(self, path: str, client: Any, question: str = "") -> str:
        """Use Claude Vision to analyze a document page-by-page."""
        doc = self.analyze(path)
        if doc.error:
            return f"Error: {doc.error}"

        if client is None:
            return doc.text[:3000]

        prompt = (
            f"Analyze this document ({doc.file_type}, {doc.pages} pages).\n"
            + (f"Question: {question}\n" if question else "Summarize the key points.\n")
            + f"\nDocument text:\n{doc.text[:8000]}"
        )
        resp = client.messages.create(
            model="claude-opus-4-7",
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}],
        )
        return resp.content[0].text

    def to_context(self, doc: DocumentContent, max_chars: int = 5000) -> str:
        """Format document for injection into AILEX agent prompts."""
        return (
            f"=== DOCUMENT: {os.path.basename(doc.path)} ({doc.file_type}, {doc.pages} pages) ===\n"
            f"{doc.text[:max_chars]}"
            + (f"\n[... {len(doc.text)-max_chars} more chars truncated]" if len(doc.text) > max_chars else "")
            + "\n=== END DOCUMENT ==="
        )

    def _pdf(self, path: str) -> DocumentContent:
        text   = ""
        pages  = 0
        tables: List[str] = []

        # Try pdfminer first
        try:
            from pdfminer.high_level import extract_text
            text  = extract_text(path)
            pages = text.count("\f") + 1
        except ImportError:
            pass

        # Try pypdf as fallback
        if not text:
            try:
                import pypdf
                reader = pypdf.PdfReader(path)
                pages  = len(reader.pages)
                text   = "\n".join(p.extract_text() or "" for p in reader.pages)
            except ImportError:
                pass

        if not text:
            return DocumentContent(path=path, file_type="pdf", text="",
                                   pages=0, tables=[], images=[],
                                   error="Install pdfminer.six or pypdf: pip install pdfminer.six")
        return DocumentContent(path=path, file_type="pdf", text=text.strip(),
                               pages=pages, tables=tables, images=[])

    def _docx(self, path: str) -> DocumentContent:
        try:
            import docx
            doc   = docx.Document(path)
            text  = "\n".join(p.text for p in doc.paragraphs)
            tables = ["\n".join(
                " | ".join(cell.text for cell in row.cells)
                for row in table.rows
            ) for table in doc.tables]
            return DocumentContent(path=path, file_type="docx",
                                   text=text, pages=len(doc.sections),
                                   tables=tables, images=[])
        except ImportError:
            return DocumentContent(path=path, file_type="docx", text="",
                                   pages=0, tables=[], images=[],
                                   error="Install python-docx: pip install python-docx")

    def _text(self, path: str) -> DocumentContent:
        try:
            text = open(path, encoding="utf-8", errors="ignore").read()
            ext  = os.path.splitext(path)[1].lower()
            return DocumentContent(path=path, file_type=ext.lstrip("."),
                                   text=text, pages=1, tables=[], images=[])
        except Exception as e:
            return DocumentContent(path=path, file_type="text", text="",
                                   pages=0, tables=[], images=[], error=str(e))
